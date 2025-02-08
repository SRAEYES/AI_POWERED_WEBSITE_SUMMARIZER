from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify
from flask_bcrypt import Bcrypt
from flask_session import Session
from pymongo import MongoClient
from docx import Document
import requests
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from transformers import pipeline
import os
import re
import concurrent.futures

app = Flask(__name__)
bcrypt = Bcrypt(app)

# Configure session
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
app.secret_key = "your_secret_key"
Session(app)

# Connect to MongoDB
client = MongoClient("mongodb://localhost:27017/")
db = client["summarizerDB"]
users_collection = db["users"]

# Initialize Hugging Face models
sentiment_analyzer = pipeline("sentiment-analysis")
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# Function to extract text from a webpage
def extract_text_from_url(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = [
            re.sub(r'\[\d+\]', '', p.text.strip())
            for p in soup.find_all("p", limit=50)
            if p.text.strip()
        ]
        text = " ".join(paragraphs)
        return text if len(text) > 50 else "Error: Extracted text too short."
    except Exception as e:
        return f"Error fetching text: {str(e)}"

# Function to analyze sentiment
def analyze_sentiment(text):
    result = sentiment_analyzer(text[:500])
    return result[0]['label']

# Function to split text into chunks
def chunk_text(text, chunk_size=800):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, current_chunk = [], ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) < chunk_size:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

# Function to summarize text
def summarize_text(text):
    text = text.strip()
    if not text:
        return "Error: No text provided for summarization."

    chunks = chunk_text(text, chunk_size=800)
    summaries = []

    def summarize_chunk(chunk):
        try:
            return summarizer(chunk, max_length=150, min_length=50, do_sample=False)[0]['summary_text']
        except Exception:
            return "Error: Summarization failed for this chunk."

    with concurrent.futures.ThreadPoolExecutor() as executor:
        summaries = list(executor.map(summarize_chunk, chunks))

    return " ".join(summaries)

# Function to translate text
def translate_text(text, target_language="hi"):
    translator = GoogleTranslator(source="auto", target=target_language)
    MAX_CHARS = 3000
    translated_parts = []

    for i in range(0, len(text), MAX_CHARS):
        chunk = text[i:i + MAX_CHARS]
        try:
            translated_parts.append(translator.translate(chunk))
        except Exception:
            translated_parts.append(f"Error translating chunk: {chunk[:50]}...")

    return " ".join(translated_parts)

# Function to create a Word document
def create_word_doc(original, summarized, translated):
    doc = Document()
    doc.add_heading("Translated Content Report", level=1)

    doc.add_heading("Original Content", level=2)
    doc.add_paragraph(original[:2000] + "...")

    if summarized:
        doc.add_heading("Summarized Content", level=2)
        doc.add_paragraph(summarized[:2000] + "...")

    doc.add_heading("Translated Content", level=2)
    doc.add_paragraph(translated[:2000] + "...")

    file_path = "static/translated_output.docx"
    doc.save(file_path)
    return file_path

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        hashed_pw = bcrypt.generate_password_hash(password).decode("utf-8")

        if users_collection.find_one({"email": email}):
            return "User already exists!"

        users_collection.insert_one({"email": email, "password": hashed_pw})
        return redirect(url_for("login"))

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]

        user = users_collection.find_one({"email": email})
        if user and bcrypt.check_password_hash(user["password"], password):
            session["user"] = email
            return redirect(url_for("dashboard"))
        else:
            return "Invalid Credentials!"

    return render_template("login.html")

@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect(url_for("login"))
    return render_template("dashboard.html", user=session["user"])

@app.route("/process", methods=["POST"])
def process():
    if "user" not in session:
        return jsonify({"error": "Unauthorized access. Please log in."}), 403

    url = request.form["url"]
    language_option = request.form["language"]
    option = request.form["option"]

    language_map = {
        "1": "ta", "2": "hi", "3": "te", "4": "or", "5": "kn",
        "6": "ml", "7": "bn", "8": "gu", "9": "mr", "10": "pa",
        "11": "ur", "12": "en"
    }
    target_language = language_map.get(language_option, "en")

    original_text = extract_text_from_url(url)
    if "Error" in original_text:
        return render_template("error.html", message=original_text)

    sentiment = analyze_sentiment(original_text)
    if sentiment == "NEGATIVE":
        return render_template("error.html", message="Harsh content detected.")

    summarized_text = summarize_text(original_text) if option == "2" else None
    text_to_translate = summarized_text if summarized_text else original_text
    translated_text = translate_text(text_to_translate, target_language)

    doc_path = create_word_doc(original_text, summarized_text, translated_text)

    return render_template("result.html", original=original_text, summarized=summarized_text, translated=translated_text, doc_path=doc_path)

@app.route("/download")
def download():
    return send_file("static/translated_output.docx", as_attachment=True)

@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("index"))

if __name__ == "__main__":
    app.run(debug=True)