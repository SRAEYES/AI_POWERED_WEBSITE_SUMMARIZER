from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify
from flask_bcrypt import Bcrypt
from flask_session import Session
from pymongo import MongoClient
from docx import Document
import requests
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from transformers import pipeline
import os

app = Flask(__name__)
bcrypt = Bcrypt(app)

# Configure session
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
app.secret_key = "your_secret_key"
Session(app)

# Connect to MongoDB
client = MongoClient("mongodb://localhost:27017/")  # Ensure MongoDB is running
db = client["summarizerDB"]
users_collection = db["users"]

# Initialize Hugging Face pipelines
sentiment_analyzer = pipeline("sentiment-analysis")
summarizer = pipeline("summarization")

# Function to extract text from a webpage
def extract_text_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = [p.text.strip() for p in soup.find_all("p") if p.text.strip()]
        return " ".join(paragraphs)
    except Exception as e:
        return f"Error fetching text: {str(e)}"

# Function to analyze sentiment
def analyze_sentiment(text):
    result = sentiment_analyzer(text[:500])  # Limiting to first 500 characters
    return result[0]['label']

# Function to translate text
def translate_text(text, target_language):
    translator = GoogleTranslator(source='auto', target=target_language)
    return translator.translate(text)

# Function to split text into chunks for summarization
def chunk_text(text, chunk_size=1024):
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

# Function to summarize content
def summarize_text(text):
    chunks = chunk_text(text)
    summarized_chunks = [summarizer(chunk, max_length=150, min_length=50, do_sample=False)[0]['summary_text'] for chunk in chunks]
    return " ".join(summarized_chunks)

# Function to create a Word document
def create_word_doc(original, summarized, translated):
    doc = Document()
    doc.add_heading("Translated Content Report", level=1)

    doc.add_heading("Original Content", level=2)
    doc.add_paragraph(original)

    if summarized:
        doc.add_heading("Summarized Content", level=2)
        doc.add_paragraph(summarized)

    doc.add_heading("Translated Content", level=2)
    doc.add_paragraph(translated)

    file_path = "static/translated_output.docx"
    doc.save(file_path)
    return file_path

# Route for home page
@app.route("/")
def index():
    return render_template("index.html")

# Route for user registration
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        hashed_pw = bcrypt.generate_password_hash(password).decode("utf-8")

        # Check if user already exists
        if users_collection.find_one({"email": email}):
            return "User already exists!"

        # Store user in MongoDB
        users_collection.insert_one({"email": email, "password": hashed_pw})
        return redirect(url_for("login"))
    
    return render_template("register.html")

# Route for user login
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]

        user = users_collection.find_one({"email": email})
        if user and bcrypt.check_password_hash(user["password"], password):
            session["user"] = email
            return redirect(url_for("dashboard"))  # ✅ FIXED: Ensure dashboard route exists
        else:
            return "Invalid Credentials!"

    return render_template("login.html")

# ✅ FIXED: Added missing `dashboard` route
@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect(url_for("login"))  # Redirect if not logged in
    return render_template("dashboard.html", user=session["user"])

# Route for summarization
@app.route("/summarize", methods=["POST"])
def summarize():
    if "user" not in session:
        return redirect(url_for("login"))  # Redirect if not logged in

    text = request.form["text"]
    summary = summarize_text(text)  # ✅ FIXED: Changed function name
    return jsonify({"summary": summary})

# Route for user logout
@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("index"))

# Route for processing summarization and translation
@app.route("/process", methods=["POST"])
def process():
    if "user" not in session:
        return jsonify({"error": "Unauthorized access. Please log in."}), 403

    url = request.form["url"]
    language_option = request.form["language"]
    option = request.form["option"]

    # Map language option to Google Translate codes
    language_map = {
        "1": "ta", "2": "hi", "3": "te", "4": "or", "5": "kn",
        "6": "ml", "7": "bn", "8": "gu", "9": "mr", "10": "pa",
        "11": "ur", "12": "en"
    }
    target_language = language_map.get(language_option, "en")

    original_text = extract_text_from_url(url)

    sentiment = analyze_sentiment(original_text)
    if sentiment == "NEGATIVE":
        return render_template("error.html", message="Harsh content detected.")

    summarized_text = summarize_text(original_text) if option == "2" else None
    text_to_translate = summarized_text if summarized_text else original_text
    translated_text = translate_text(text_to_translate, target_language)

    doc_path = create_word_doc(original_text, summarized_text, translated_text)

    return render_template("result.html", original=original_text, summarized=summarized_text, translated=translated_text, doc_path=doc_path)

# Route for downloading the translated document
@app.route("/download")
def download():
    return send_file("static/translated_output.docx", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
